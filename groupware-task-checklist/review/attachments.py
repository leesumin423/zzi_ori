"""첨부파일에서 텍스트(또는 스캔본이면 페이지 이미지)를 뽑아낸다.

지원:
  - PDF: 텍스트 레이어가 있으면 텍스트 추출. 텍스트가 거의 없으면(=스캔본으로 추정)
    페이지를 이미지로 렌더링해서 Claude Vision에 그대로 넘긴다 (poppler 등 외부 프로그램 불필요).
  - DOCX / XLSX: 텍스트 추출.
미지원 (수동 확인 필요, 리포트에 안내만 표시):
  - HWP(한글) — 마땅한 순수 파이썬 라이브러리가 없어 v1에서는 처리하지 않음.
"""

from __future__ import annotations

from pathlib import Path
from typing import List

from .models import Attachment

_SCANNED_TEXT_THRESHOLD = 40  # 페이지당 이 글자 수보다 적으면 스캔본으로 간주


def _extract_pdf(path: Path, max_pages: int) -> Attachment:
    import fitz  # PyMuPDF

    doc = fitz.open(path)
    texts: List[str] = []
    for page in doc:
        texts.append(page.get_text())
    full_text = "\n".join(texts).strip()

    avg_chars_per_page = len(full_text) / max(len(doc), 1)
    if avg_chars_per_page >= _SCANNED_TEXT_THRESHOLD:
        doc.close()
        return Attachment(filename=path.name, text=full_text)

    # 스캔본으로 추정 → 페이지 이미지 렌더링 (최대 max_pages 페이지)
    images: List[bytes] = []
    page_count = len(doc)
    for i, page in enumerate(doc):
        if i >= max_pages:
            break
        pix = page.get_pixmap(dpi=150)
        images.append(pix.tobytes("png"))
    doc.close()

    note = ""
    if page_count > max_pages:
        note = f"{page_count}페이지 중 {max_pages}페이지만 검토에 포함됨"

    return Attachment(filename=path.name, text=full_text, image_pngs=images, note=note)


def _extract_docx(path: Path) -> Attachment:
    import docx

    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return Attachment(filename=path.name, text="\n".join(parts))


def _extract_xlsx(path: Path) -> Attachment:
    import openpyxl

    wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"[시트: {ws.title}]")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                parts.append(" | ".join(cells))
    wb.close()
    return Attachment(filename=path.name, text="\n".join(parts))


def extract_attachment(path: Path, max_pages: int = 8) -> Attachment:
    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            return _extract_pdf(path, max_pages)
        if suffix == ".docx":
            return _extract_docx(path)
        if suffix in (".xlsx", ".xlsm"):
            return _extract_xlsx(path)
        if suffix == ".hwp":
            return Attachment(
                filename=path.name, unsupported=True,
                note="HWP 파일은 자동으로 내용을 읽을 수 없습니다. 수동으로 확인해주세요.",
            )
        if suffix in (".jpg", ".jpeg", ".png"):
            return Attachment(filename=path.name, image_pngs=[path.read_bytes()])
    except Exception as e:  # noqa: BLE001 - 첨부파일 하나 실패해도 나머지는 계속 진행
        return Attachment(filename=path.name, unsupported=True, note=f"읽기 실패: {e}")

    return Attachment(
        filename=path.name, unsupported=True,
        note=f"지원하지 않는 파일 형식({suffix})입니다. 수동으로 확인해주세요.",
    )
